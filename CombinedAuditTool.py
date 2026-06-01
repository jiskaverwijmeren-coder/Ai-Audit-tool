import io
import os
import tempfile
import time
from typing import Dict, List, Tuple

import nltk
import numpy as np
import pandas as pd
import PyPDF2
import streamlit as st
import whisper
from docx import Document
from dotenv import load_dotenv
from mistralai import Mistral
from nltk.tokenize import sent_tokenize
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
from streamlit_mic_recorder import mic_recorder
from supabase import create_client

from analyzer import analyseer_iso


st.set_page_config(page_title="AI Audit Suite", page_icon="🔍", layout="wide")

load_dotenv()

SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")
supabase_client = create_client(SUPABASE_URL, SUPABASE_KEY)

PROBLEM_SIGNALS = [
    "missing documentation",
    "delay or overdue",
    "risk or exposure",
    "unclear ownership",
    "miscommunication",
    "outdated information",
    "duplicate or redundant work",
    "bottleneck or blockage",
    "non-compliance",
    "lack of controls",
    "unauthorized access",
    "inconsistency",
    "error or mistake",
    "incomplete process",
]

MAX_TOKENS_PER_DOCUMENT = 25_000


@st.cache_resource
def load_whisper_model():
    return whisper.load_model("base")


@st.cache_resource
def setup_nltk():
    nltk.download("punkt", quiet=True)
    nltk.download("punkt_tab", quiet=True)
    return True


@st.cache_resource
def load_embedding_model():
    return SentenceTransformer("all-MiniLM-L6-v2")


whisper_model = load_whisper_model()
setup_nltk()
embed_model = load_embedding_model()


def init_session_state() -> None:
    defaults = {
        "chat_geschiedenis": {},
        "document_list": [],
        "live_audit_log": [],
        "all_results": {},
        "feedback_store": None,
    }

    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

    if st.session_state.feedback_store is None:
        st.session_state.feedback_store = laad_feedback_supabase()


def laad_feedback_supabase() -> Dict[str, List[dict]]:
    try:
        goede_feedback = (
            supabase_client.table("audit_feedback")
            .select("*")
            .eq("positief", True)
            .order("timestamp", desc=True)
            .limit(20)
            .execute()
        )

        slechte_feedback = (
            supabase_client.table("audit_feedback")
            .select("*")
            .eq("positief", False)
            .order("timestamp", desc=True)
            .limit(20)
            .execute()
        )

        return {
            "goede_vragen": [
                {
                    "zin": record["zin"],
                    "issue_type": record["issue_type"],
                    "vragen": record["vragen"],
                    "score": record["score"],
                }
                for record in goede_feedback.data
            ],
            "slechte_vragen": [
                {
                    "zin": record["zin"],
                    "issue_type": record["issue_type"],
                    "vragen": record["vragen"],
                    "score": record["score"],
                }
                for record in slechte_feedback.data
            ],
        }
    except Exception as exc:
        st.warning(f"Feedback kon niet worden geladen: {exc}")
        return {"goede_vragen": [], "slechte_vragen": []}


def sla_feedback_op_supabase(
    zin: str,
    issue_type: str,
    vragen: str,
    score: float,
    positief: bool,
) -> bool:
    try:
        supabase_client.table("audit_feedback").insert(
            {
                "zin": zin,
                "issue_type": issue_type,
                "vragen": vragen,
                "score": score,
                "positief": positief,
            }
        ).execute()
        return True
    except Exception as exc:
        st.warning(f"Feedback kon niet worden opgeslagen: {exc}")
        return False


def verwijder_feedback_supabase(zin: str, issue_type: str) -> bool:
    try:
        (
            supabase_client.table("audit_feedback")
            .delete()
            .eq("zin", zin)
            .eq("issue_type", issue_type)
            .execute()
        )
        return True
    except Exception as exc:
        st.warning(f"Feedback kon niet worden verwijderd: {exc}")
        return False


def apply_css() -> None:
    st.markdown(
        """
        <style>
            .stDeployButton,
            #MainMenu,
            footer,
            [data-testid="stHeader"] {
                display: none !important;
            }

            [data-testid="stSidebarNav"] button,
            button[kind="headerNoSpacing"],
            [data-testid="stSidebarCollapseButton"] {
                display: none !important;
            }

            [data-testid="stSidebar"] {
                min-width: 300px !important;
                max-width: 300px !important;
            }

            .stApp {
                background-color: #f8f9fa;
                color: #1a1a1a;
            }

            section[data-testid="stSidebar"] {
                background-color: #005B94;
                border-right: 2px solid #00AEEF;
            }

            .stButton > button {
                background-color: #005B94;
                color: white;
                border: none;
                border-radius: 8px;
                font-weight: 600;
            }

            .stButton > button:hover {
                background-color: #00AEEF;
                color: white;
            }

            [data-testid="stMetric"] {
                background-color: #6AAA3A;
                border: 1px solid #00AEEF;
                border-radius: 10px;
                padding: 12px;
                color: white;
            }

            [data-testid="stExpander"] {
                border: 1px solid #005B94;
                border-radius: 8px;
                background-color: #f0f8ff;
            }

            h1,
            h2,
            h3 {
                color: #005B94;
            }

            hr {
                border-color: #00AEEF;
            }

            [data-testid="stChatInput"] {
                background-color: #f0f8ff !important;
                border: 1px solid #00AEEF !important;
                border-radius: 8px !important;
            }

            [data-testid="stChatInput"] textarea {
                background-color: transparent !important;
                color: #1a1a1a !important;
            }

            [data-testid="stBottom"] > div {
                background-color: #f8f9fa !important;
            }

            .doc-card {
                background-color: #f0f8ff;
                border: 1px solid #00AEEF;
                border-radius: 10px;
                color: #1a1a1a;
                display: flex;
                align-items: center;
                justify-content: space-between;
                font-size: 0.9rem;
                margin-bottom: 8px;
                padding: 10px 16px;
            }

            .transcription-box {
                background-color: #f0f8ff;
                border: 1px solid #00AEEF;
                border-radius: 10px;
                color: #1a1a1a;
                font-size: 0.92rem;
                line-height: 1.6;
                margin: 10px 0;
                padding: 14px 18px;
                white-space: pre-wrap;
            }
        </style>
        """,
        unsafe_allow_html=True,
    )


def extract_text_from_pdf(file) -> str:
    reader = PyPDF2.PdfReader(file)
    text_parts = []

    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)

    return "\n".join(text_parts)


def transcribe_audio_file(file_bytes: bytes, suffix: str) -> str:
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as temp_file:
        temp_file.write(file_bytes)
        temp_path = temp_file.name

    try:
        result = whisper_model.transcribe(temp_path, language="nl")
        return result["text"]
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def split_into_sentences(text: str) -> List[str]:
    sentences = sent_tokenize(text, language="dutch")
    return [sentence.strip() for sentence in sentences if len(sentence.strip()) > 30]


def detect_problem_sentences(sentences: List[str], threshold: float = 0.30) -> List[dict]:
    if not sentences:
        return []

    signal_embeddings = embed_model.encode(PROBLEM_SIGNALS)
    sentence_embeddings = embed_model.encode(sentences)
    similarity_matrix = cosine_similarity(sentence_embeddings, signal_embeddings)

    results = []
    for index, scores in enumerate(similarity_matrix):
        best_index = int(np.argmax(scores))
        best_score = float(scores[best_index])

        if best_score >= threshold:
            results.append(
                {
                    "sentence": sentences[index],
                    "issue_type": PROBLEM_SIGNALS[best_index],
                    "score": best_score,
                }
            )

    return sorted(results, key=lambda item: item["score"], reverse=True)


def generate_audit_questions(client: Mistral, sentence: str, issue_type: str) -> str:
    feedback = st.session_state.feedback_store
    few_shot_examples = ""

    goede_voorbeelden = [
        entry
        for entry in feedback["goede_vragen"]
        if entry["issue_type"] == issue_type
    ][-3:]

    slechte_voorbeelden = [
        entry
        for entry in feedback["slechte_vragen"]
        if entry["issue_type"] == issue_type
    ][-2:]

    if goede_voorbeelden:
        few_shot_examples += "\n\nVoorbeelden van passende auditvragen:\n"
        for example in goede_voorbeelden:
            few_shot_examples += (
                f'Zin: "{example["zin"]}"\n'
                f'Vragen:\n{example["vragen"]}\n\n'
            )

    if slechte_voorbeelden:
        few_shot_examples += "\nVoorbeelden van minder passende auditvragen:\n"
        for example in slechte_voorbeelden:
            few_shot_examples += (
                f'Zin: "{example["zin"]}"\n'
                f'Vragen:\n{example["vragen"]}\n\n'
            )

    prompt = f"""You are a senior internal auditor reviewing a document.

The following sentence was flagged as potentially problematic:
"{sentence}"

Detected issue category: {issue_type}
{few_shot_examples}

Generate exactly 3 sharp, professional audit follow-up questions.
- Be specific to the sentence content.
- Do not repeat the sentence.
- Return only the 3 questions as a numbered list.
"""

    response = client.chat.complete(
        model="mistral-large-latest",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3,
        max_tokens=300,
    )
    return response.choices[0].message.content.strip()


def haal_gecachete_analyse_op(
    api_sleutel: str,
    tekst: str,
    geselecteerde_normen: Tuple[str, ...],
):
    client = Mistral(api_key=api_sleutel)
    max_retries = 5

    for attempt in range(max_retries):
        try:
            return analyseer_iso(client, tekst, normen=list(geselecteerde_normen))
        except Exception as exc:
            if "429" in str(exc) and attempt < max_retries - 1:
                wait_time = 10 * (attempt + 1)
                st.warning(
                    f"Rate limit bereikt. Nieuwe poging over {wait_time} seconden "
                    f"({attempt + 1}/{max_retries})."
                )
                time.sleep(wait_time)
            else:
                raise exc


def genereer_word_rapport(
    doc_name: str,
    data: dict,
    context: str,
    gefilterde_bevindingen: List[dict],
) -> bytes:
    document = Document()
    document.add_heading(f"ISO Rapport — {doc_name}", 0)

    if context:
        document.add_heading("Projectcontext", level=1)
        document.add_paragraph(context)

    document.add_heading("Managementsamenvatting", level=1)
    document.add_paragraph(data["samenvatting"])

    document.add_heading(f"Bevindingen ({len(gefilterde_bevindingen)})", level=1)
    for bevinding in gefilterde_bevindingen:
        document.add_heading(
            f"{bevinding['norm']} - {bevinding['clausule']}: {bevinding['titel']}",
            level=2,
        )
        document.add_paragraph(f"Ernst: {bevinding['ernst'].capitalize()}")
        document.add_paragraph(f"Probleem: {bevinding['beschrijving']}")
        document.add_paragraph(f"Aanbeveling: {bevinding['aanbeveling']}")

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def genereer_totaal_word_rapport(all_results: dict, context: str) -> bytes:
    document = Document()
    document.add_heading("ISO Totaalrapport — Alle documenten", 0)

    if context:
        document.add_heading("Projectcontext", level=1)
        document.add_paragraph(context)

    for doc_name, result in all_results.items():
        data = result["iso_data"]

        document.add_heading(f"Document: {doc_name}", level=1)
        document.add_heading("Managementsamenvatting", level=2)
        document.add_paragraph(data["samenvatting"])

        document.add_heading(f"Bevindingen ({len(data['bevindingen'])})", level=2)
        for bevinding in data["bevindingen"]:
            document.add_heading(
                f"{bevinding['norm']} - {bevinding['clausule']}: {bevinding['titel']}",
                level=3,
            )
            document.add_paragraph(f"Ernst: {bevinding['ernst'].capitalize()}")
            document.add_paragraph(f"Probleem: {bevinding['beschrijving']}")
            document.add_paragraph(f"Aanbeveling: {bevinding['aanbeveling']}")

        document.add_page_break()

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def genereer_vragen_word_rapport(all_results: dict) -> bytes:
    document = Document()
    document.add_heading("Auditopvolgingsvragen — Alle documenten", 0)

    for doc_name, result in all_results.items():
        risk_results = result["risk_results"]
        if not risk_results:
            continue

        document.add_heading(f"Document: {doc_name}", level=1)
        for index, item in enumerate(risk_results, 1):
            document.add_heading(
                f"Vraagset {index} — {item['issue_type'].upper()} "
                f"(score: {item['score']:.2f})",
                level=2,
            )
            document.add_paragraph(f"Gemarkeerde zin:\n{item['sentence']}")
            document.add_heading("Auditopvolgingsvragen", level=3)
            document.add_paragraph(item["questions"])

        document.add_page_break()

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def render_sidebar():
    with st.sidebar:
        st.header("Instellingen")

        api_key = st.text_input(
            "Mistral API-sleutel",
            type="password",
            value=os.getenv("MISTRAL_API_KEY", ""),
            help="Gebruik de API-sleutel van Mistral.",
        )

        st.divider()
        st.subheader("Projectcontext")
        project_context = st.text_input(
            "Beschrijf het project (optioneel)",
            help="Bijvoorbeeld: aanleg snelweg A12.",
        )

        st.divider()
        st.subheader("ISO-normen")
        enable_9001 = st.toggle("ISO 9001 (Kwaliteit)", value=True)
        enable_14001 = st.toggle("ISO 14001 (Milieu)", value=True)
        enable_45001 = st.toggle("ISO 45001 (Veiligheid)", value=True)

        st.divider()
        st.subheader("Risk scanner")
        threshold = st.slider(
            "Detectiegevoeligheid",
            min_value=0.20,
            max_value=0.60,
            value=0.30,
            step=0.05,
        )

        st.divider()
        st.subheader("Live auditmodus")
        enable_audit_mode = st.toggle("Schakel live auditmodus in", value=False)

        st.divider()
        st.caption("AI Audit Suite · v2.1")

    return (
        api_key,
        project_context,
        enable_9001,
        enable_14001,
        enable_45001,
        threshold,
        enable_audit_mode,
    )


def render_header() -> None:
    col1, col2 = st.columns([10, 9])
    with col2:
        st.image("logo.png", width=200)

    st.title("🔍 AI Audit Suite")


def render_audio_upload() -> None:
    st.subheader("Audiobestand uploaden")
    st.caption("Upload een MP3-, WAV- of M4A-bestand. Het bestand wordt automatisch getranscribeerd.")

    audio_file = st.file_uploader(
        "Kies een audiobestand",
        type=["mp3", "wav", "m4a"],
        help="Ondersteunde formaten: MP3, WAV en M4A.",
        key="audio_uploader",
    )

    if audio_file is None:
        return

    st.audio(audio_file, format=f"audio/{audio_file.name.split('.')[-1]}")

    if st.button("Transcribeer audiobestand"):
        suffix = "." + audio_file.name.split(".")[-1]
        with st.spinner(f"Transcriberen van '{audio_file.name}'..."):
            transcript = transcribe_audio_file(audio_file.read(), suffix)
            st.session_state[f"audio_transcript_{audio_file.name}"] = transcript

    transcript_key = f"audio_transcript_{audio_file.name}"
    if transcript_key not in st.session_state:
        return

    transcript = st.session_state[transcript_key]
    st.success("Transcriptie voltooid.")
    st.markdown("**Getranscribeerde tekst:**")
    st.markdown(f'<div class="transcription-box">{transcript}</div>', unsafe_allow_html=True)

    col1, col2 = st.columns(2)
    with col1:
        doc_name_audio = st.text_input(
            "Naam voor dit document:",
            value=audio_file.name.rsplit(".", 1)[0],
            key=f"name_{audio_file.name}",
        )

    with col2:
        st.write("")
        st.write("")
        if st.button("Voeg toe aan documentenlijst", key=f"add_{audio_file.name}"):
            existing_names = [doc["name"] for doc in st.session_state.document_list]
            name_to_use = f"Audio - {doc_name_audio}"

            if name_to_use not in existing_names:
                st.session_state.document_list.append(
                    {"name": name_to_use, "text": transcript}
                )
                st.success(f"'{name_to_use}' is toegevoegd aan de documentenlijst.")
            else:
                st.warning(f"'{name_to_use}' staat al in de documentenlijst.")

            st.rerun()


def render_document_upload() -> None:
    st.divider()
    st.subheader("Documenten uploaden")

    uploaded_files = st.file_uploader(
        "Kies één of meerdere bestanden",
        type=["txt", "pdf"],
        help="Ondersteunde bestandstypen: TXT en PDF.",
        accept_multiple_files=True,
        key="doc_uploader",
    )

    if not uploaded_files:
        return

    existing_names = [doc["name"] for doc in st.session_state.document_list]
    added = 0

    for uploaded_file in uploaded_files:
        if uploaded_file.name in existing_names:
            continue

        if uploaded_file.type == "application/pdf":
            text = extract_text_from_pdf(uploaded_file)
        else:
            text = uploaded_file.read().decode("utf-8")

        if text.strip():
            st.session_state.document_list.append(
                {"name": uploaded_file.name, "text": text}
            )
            added += 1

    if added:
        st.success(f"{added} nieuw(e) document(en) toegevoegd.")
        st.rerun()


def render_document_list() -> None:
    st.divider()
    st.subheader("Documentenlijst")

    if not st.session_state.document_list:
        st.info("Nog geen documenten toegevoegd.")
        return

    st.caption(f"{len(st.session_state.document_list)} document(en) klaar voor analyse")
    document_to_remove = None

    for index, document in enumerate(st.session_state.document_list):
        col1, col2 = st.columns([6, 1])
        with col1:
            st.markdown(
                (
                    f'<div class="doc-card">📄 <strong>{document["name"]}</strong>'
                    f' &nbsp;·&nbsp; {len(document["text"]):,} tekens</div>'
                ),
                unsafe_allow_html=True,
            )
        with col2:
            if st.button("Verwijder", key=f"remove_{index}", help=f"Verwijder {document['name']}"):
                document_to_remove = index

    if document_to_remove is not None:
        st.session_state.document_list.pop(document_to_remove)
        st.rerun()

    if st.button("Verwijder alle documenten", type="secondary"):
        st.session_state.document_list = []
        st.session_state.all_results = {}
        st.session_state.chat_geschiedenis = {}
        st.rerun()


def get_active_norms(
    enable_9001: bool,
    enable_14001: bool,
    enable_45001: bool,
) -> Tuple[str, ...]:
    return tuple(
        norm
        for norm, enabled in [
            ("ISO 9001", enable_9001),
            ("ISO 14001", enable_14001),
            ("ISO 45001", enable_45001),
        ]
        if enabled
    )


def validate_analysis_inputs(
    api_key: str,
    enable_9001: bool,
    enable_14001: bool,
    enable_45001: bool,
) -> bool:
    analysis_blocked = False

    if not st.session_state.document_list:
        st.info("Voeg documenten toe om de analyse te starten.")
        analysis_blocked = True

    if not api_key:
        st.warning("Voer je Mistral API-sleutel in via de sidebar.")
        analysis_blocked = True

    if not (enable_9001 or enable_14001 or enable_45001):
        st.warning("Selecteer minimaal één ISO-norm in de sidebar.")
        analysis_blocked = True

    return not analysis_blocked


def run_document_analysis(
    api_key: str,
    project_context: str,
    active_norms: Tuple[str, ...],
    threshold: float,
) -> None:
    client = Mistral(api_key=api_key)
    all_results = {}

    for document in st.session_state.document_list:
        doc_name = document["name"]
        raw_text = document["text"]

        st.markdown(f"### Bezig met: **{doc_name}**")

        estimated_tokens = len(raw_text) // 4
        if estimated_tokens > MAX_TOKENS_PER_DOCUMENT:
            st.error(
                f"'{doc_name}' is te groot "
                f"({estimated_tokens} tokens, maximaal {MAX_TOKENS_PER_DOCUMENT})."
            )
            continue

        text_for_analysis = (
            f"CONTEXT:\n{project_context}\n\nDOCUMENT:\n{raw_text}"
            if project_context
            else raw_text
        )

        with st.spinner(f"Risk scanner: {doc_name}"):
            sentences = split_into_sentences(raw_text)
            detected = detect_problem_sentences(sentences, threshold=threshold)

        risk_results = []
        if detected:
            progress = st.progress(0, text="Auditvragen genereren...")
            for index, item in enumerate(detected):
                progress.progress(
                    index / len(detected),
                    text=f"Vraag {index + 1}/{len(detected)} — {doc_name}",
                )

                for attempt in range(4):
                    try:
                        questions = generate_audit_questions(
                            client,
                            item["sentence"],
                            item["issue_type"],
                        )
                        risk_results.append({**item, "questions": questions})
                        break
                    except Exception as exc:
                        if "429" in str(exc) and attempt < 3:
                            time.sleep(2**attempt)
                        else:
                            st.error(f"Fout bij zin {index + 1}: {exc}")
                            break

                if index < len(detected) - 1:
                    time.sleep(3)

                progress.progress((index + 1) / len(detected))

            progress.empty()

        time.sleep(10)
        with st.spinner(f"ISO-analyse: {doc_name}"):
            iso_data = haal_gecachete_analyse_op(api_key, text_for_analysis, active_norms)

        all_results[doc_name] = {
            "raw_text": raw_text,
            "tekst_voor_analyse": text_for_analysis,
            "sentences_count": len(sentences),
            "risk_results": risk_results,
            "iso_data": iso_data,
        }

        st.success(f"{doc_name} is geanalyseerd.")

    st.session_state.all_results = all_results
    st.session_state.chat_geschiedenis = {}
    st.rerun()


def render_results(api_key: str, project_context: str) -> None:
    if not st.session_state.all_results:
        return

    st.divider()
    st.header("Analyseresultaten")

    if len(st.session_state.all_results) > 1:
        st.subheader("Totaalrapporten")
        col_total_1, col_total_2 = st.columns(2)

        with col_total_1:
            total_word = genereer_totaal_word_rapport(
                st.session_state.all_results,
                project_context,
            )
            st.download_button(
                "Download ISO-totaalrapport (Word)",
                data=total_word,
                file_name="iso_totaalrapport.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="dl_totaal_iso",
            )

        with col_total_2:
            questions_word = genereer_vragen_word_rapport(st.session_state.all_results)
            st.download_button(
                "Download vragenrapport (Word)",
                data=questions_word,
                file_name="auditopvolgingsvragen.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="dl_totaal_vragen",
            )

        st.divider()

    for doc_name, result in st.session_state.all_results.items():
        with st.expander(f"Document: {doc_name}", expanded=True):
            render_risk_results(doc_name, result)
            st.divider()
            render_iso_results(doc_name, result, project_context)
            st.divider()
            render_document_chat(api_key, doc_name, result)


def render_risk_results(doc_name: str, result: dict) -> None:
    st.subheader("Risk signal scanner")
    st.caption(f"{result['sentences_count']} zinnen geanalyseerd")

    risk_results = result["risk_results"]

    if not risk_results:
        st.warning("Geen risicosignalen gevonden bij de huidige gevoeligheidsdrempel.")
        return

    st.success(f"{len(risk_results)} zin(nen) gemarkeerd")

    for index, item in enumerate(risk_results, 1):
        with st.expander(
            f"Vraagset {index} — {item['issue_type'].upper()} "
            f"(score: {item['score']:.2f})"
        ):
            st.markdown("**Gemarkeerde zin:**")
            st.markdown(f"> {item['sentence']}")
            st.markdown("**Auditopvolgingsvragen:**")
            st.markdown(item["questions"])
            st.caption("Feedback wordt opgeslagen in Supabase.")

            col_positive, col_negative = st.columns([1, 1])
            with col_positive:
                if st.button("Goede vragen", key=f"pos_{doc_name}_{index}"):
                    success = sla_feedback_op_supabase(
                        zin=item["sentence"],
                        issue_type=item["issue_type"],
                        vragen=item["questions"],
                        score=item["score"],
                        positief=True,
                    )
                    if success:
                        st.session_state.feedback_store = laad_feedback_supabase()
                        st.success("Feedback opgeslagen.")

            with col_negative:
                if st.button("Slechte vragen", key=f"neg_{doc_name}_{index}"):
                    success = sla_feedback_op_supabase(
                        zin=item["sentence"],
                        issue_type=item["issue_type"],
                        vragen=item["questions"],
                        score=item["score"],
                        positief=False,
                    )
                    if success:
                        st.session_state.feedback_store = laad_feedback_supabase()
                        st.warning("Feedback opgeslagen.")


def render_iso_results(doc_name: str, result: dict, project_context: str) -> None:
    st.subheader("ISO-analyse")

    data = result["iso_data"]
    st.info(data["samenvatting"])

    high = sum(1 for item in data["bevindingen"] if item["ernst"] == "hoog")
    medium = sum(1 for item in data["bevindingen"] if item["ernst"] == "gemiddeld")
    low = sum(1 for item in data["bevindingen"] if item["ernst"] == "laag")

    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Totaal", len(data["bevindingen"]))
    col2.metric("Hoog", high)
    col3.metric("Gemiddeld", medium)
    col4.metric("Laag", low)

    st.subheader("Filters")
    col_filter_1, col_filter_2 = st.columns(2)

    with col_filter_1:
        filter_ernst = st.multiselect(
            "Filter op ernst",
            ["hoog", "gemiddeld", "laag"],
            default=["hoog", "gemiddeld", "laag"],
            key=f"ernst_{doc_name}",
        )

    with col_filter_2:
        current_norms = list({item["norm"] for item in data["bevindingen"]})
        filter_norm = st.multiselect(
            "Filter op norm",
            current_norms,
            default=current_norms,
            key=f"norm_{doc_name}",
        )

    filtered_findings = [
        item
        for item in data["bevindingen"]
        if item["ernst"] in filter_ernst and item["norm"] in filter_norm
    ]

    if not filtered_findings:
        st.warning("Geen bevindingen gevonden voor de geselecteerde filters.")
        return

    st.success(f"{len(filtered_findings)} bevinding(en) gevonden")
    col_download_1, col_download_2 = st.columns(2)

    with col_download_1:
        csv_data = pd.DataFrame(filtered_findings).to_csv(index=False).encode("utf-8")
        st.download_button(
            "Download CSV",
            data=csv_data,
            file_name=f"rapport_{doc_name}.csv",
            mime="text/csv",
            use_container_width=True,
            key=f"csv_{doc_name}",
        )

    with col_download_2:
        word_data = genereer_word_rapport(
            doc_name,
            data,
            project_context,
            filtered_findings,
        )
        st.download_button(
            "Download Word",
            data=word_data,
            file_name=f"rapport_{doc_name}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key=f"word_{doc_name}",
        )

    st.write("")
    for finding in filtered_findings:
        severity_icon = {
            "hoog": "🔴",
            "gemiddeld": "🟠",
            "laag": "🟡",
        }.get(finding["ernst"], "⚪")

        with st.expander(
            f"{severity_icon} [{finding['norm']} | {finding['clausule']}] "
            f"{finding['titel']}"
        ):
            st.markdown(f"**Ernst:** {finding['ernst'].capitalize()}")
            st.markdown(f"**Probleem:** {finding['beschrijving']}")
            st.markdown(f"**Aanbeveling:** {finding['aanbeveling']}")


def render_document_chat(api_key: str, doc_name: str, result: dict) -> None:
    st.subheader(f"Chat over: {doc_name}")
    st.caption("Stel vragen over dit specifieke document.")

    if doc_name not in st.session_state.chat_geschiedenis:
        st.session_state.chat_geschiedenis[doc_name] = []

    for message in st.session_state.chat_geschiedenis[doc_name]:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    prompt = st.chat_input(f"Vraag over {doc_name}...", key=f"chat_{doc_name}")
    if not prompt:
        return

    st.session_state.chat_geschiedenis[doc_name].append(
        {"role": "user", "content": prompt}
    )

    with st.chat_message("user"):
        st.markdown(prompt)

    with st.chat_message("assistant"):
        with st.spinner("Antwoord wordt gegenereerd..."):
            try:
                client = Mistral(api_key=api_key)
                messages = [
                    {
                        "role": "system",
                        "content": (
                            "Je bent een ISO-auditorassistent. "
                            "Beantwoord vragen uitsluitend op basis van dit document:\n\n"
                            f"{result['tekst_voor_analyse']}"
                        ),
                    }
                ]
                messages.extend(st.session_state.chat_geschiedenis[doc_name])

                chat_response = client.chat.complete(
                    model="mistral-large-latest",
                    messages=messages,
                )
                answer = chat_response.choices[0].message.content

                st.markdown(answer)
                st.session_state.chat_geschiedenis[doc_name].append(
                    {"role": "assistant", "content": answer}
                )
            except Exception as exc:
                st.error(f"Fout tijdens chatten: {exc}")


def render_document_analysis_tab(
    api_key: str,
    project_context: str,
    enable_9001: bool,
    enable_14001: bool,
    enable_45001: bool,
    threshold: float,
) -> None:
    st.caption(
        "Upload documenten of audiobestanden, detecteer risico's en genereer "
        "auditbevindingen en ISO-analyses."
    )

    render_audio_upload()
    render_document_upload()
    render_document_list()

    if not validate_analysis_inputs(api_key, enable_9001, enable_14001, enable_45001):
        return

    active_norms = get_active_norms(enable_9001, enable_14001, enable_45001)

    st.divider()
    if st.button("Analyseer alle documenten", type="primary"):
        run_document_analysis(api_key, project_context, active_norms, threshold)

    render_results(api_key, project_context)


def render_live_audit_tab(
    api_key: str,
    enable_9001: bool,
    enable_14001: bool,
    enable_45001: bool,
    enable_audit_mode: bool,
) -> None:
    st.subheader("Live auditmodus")
    st.caption(
        "Neem een auditgesprek op via de microfoon. De opname wordt "
        "getranscribeerd en geanalyseerd."
    )

    if not api_key:
        st.warning("Voer je Mistral API-sleutel in via de sidebar.")
    if not (enable_9001 or enable_14001 or enable_45001):
        st.warning("Selecteer minimaal één ISO-norm in de sidebar.")
    if not enable_audit_mode:
        st.info("Schakel live auditmodus in via de sidebar om automatische analyse te starten.")

    st.divider()

    audio_data = mic_recorder(
        start_prompt="Start opname",
        stop_prompt="Stop opname",
        key="recorder_tab2",
    )

    if audio_data is not None and len(audio_data.get("bytes", b"")) > 0:
        process_live_recording(
            api_key,
            enable_9001,
            enable_14001,
            enable_45001,
            enable_audit_mode,
            audio_data,
        )

    render_live_audit_log()


def process_live_recording(
    api_key: str,
    enable_9001: bool,
    enable_14001: bool,
    enable_45001: bool,
    enable_audit_mode: bool,
    audio_data: dict,
) -> None:
    with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as temp_audio:
        temp_audio.write(audio_data["bytes"])
        temp_audio_path = temp_audio.name

    with st.spinner("Transcriberen..."):
        result = whisper_model.transcribe(temp_audio_path, language="nl")
        st.session_state.transcribed_text = result["text"]

    if os.path.exists(temp_audio_path):
        os.remove(temp_audio_path)

    st.success("Transcriptie voltooid.")
    st.markdown("**Getranscribeerde tekst:**")
    st.markdown(
        f'<div class="transcription-box">{st.session_state.transcribed_text}</div>',
        unsafe_allow_html=True,
    )

    if not (enable_audit_mode and api_key and (enable_9001 or enable_14001 or enable_45001)):
        return

    with st.spinner("Live auditanalyse uitvoeren..."):
        client = Mistral(api_key=api_key)

        summary_prompt = (
            "Vat het volgende antwoord van de auditee samen in 1-2 zinnen:\n"
            f'"{st.session_state.transcribed_text}"'
        )
        summary_response = client.chat.complete(
            model="mistral-large-latest",
            messages=[{"role": "user", "content": summary_prompt}],
            temperature=0.2,
            max_tokens=150,
        )
        summary = summary_response.choices[0].message.content.strip()

        sentences = split_into_sentences(st.session_state.transcribed_text)
        risk_results = (
            detect_problem_sentences(sentences, threshold=0.30)
            if sentences
            else []
        )

        follow_up_questions = []
        for item in risk_results:
            questions = generate_audit_questions(
                client,
                item["sentence"],
                item["issue_type"],
            )
            follow_up_questions.append(questions)

        active_norms = get_active_norms(enable_9001, enable_14001, enable_45001)
        iso_data = haal_gecachete_analyse_op(
            api_key,
            st.session_state.transcribed_text,
            active_norms,
        )

        live_result = {
            "timestamp": time.strftime("%H:%M:%S"),
            "transcriptie": st.session_state.transcribed_text,
            "summary": summary,
            "risk_results": risk_results,
            "follow_up_questions": follow_up_questions,
            "iso_data": iso_data,
        }

        st.session_state.live_audit_results = live_result
        st.session_state.live_audit_log.append(live_result)

    render_live_audit_results()


def render_live_audit_results() -> None:
    st.divider()
    st.subheader("Live auditanalyse")

    st.markdown("### Samenvatting")
    st.info(st.session_state.live_audit_results["summary"])

    st.markdown("### Risicosignalen en opvolgingsvragen")
    if st.session_state.live_audit_results["risk_results"]:
        for index, risk in enumerate(st.session_state.live_audit_results["risk_results"], 1):
            st.markdown(f"**{index}. {risk['issue_type'].upper()}** (score: {risk['score']:.2f})")
            st.markdown(f"> {risk['sentence']}")

            if index <= len(st.session_state.live_audit_results["follow_up_questions"]):
                st.markdown("**Voorgestelde opvolgvragen:**")
                st.markdown(st.session_state.live_audit_results["follow_up_questions"][index - 1])

            st.write("")
    else:
        st.success("Geen risicosignalen gedetecteerd in deze opname.")

    st.markdown("### Gekoppelde ISO-clausules")
    if st.session_state.live_audit_results["iso_data"]["bevindingen"]:
        for finding in st.session_state.live_audit_results["iso_data"]["bevindingen"]:
            severity_icon = {
                "hoog": "🔴",
                "gemiddeld": "🟠",
                "laag": "🟡",
            }.get(finding["ernst"], "⚪")
            st.markdown(
                f"- {severity_icon} **[{finding['norm']} | {finding['clausule']}] "
                f"{finding['titel']}**"
            )
    else:
        st.info("Geen ISO-clausules gekoppeld aan deze transcriptie.")

    col_clear, col_add = st.columns(2)
    with col_clear:
        if st.button("Wis live auditresultaten"):
            del st.session_state.live_audit_results
            st.rerun()

    with col_add:
        if st.button("Voeg opname toe aan documentenlijst"):
            name = "Live opname"
            existing_names = [document["name"] for document in st.session_state.document_list]

            if name not in existing_names:
                st.session_state.document_list.append(
                    {"name": name, "text": st.session_state.transcribed_text}
                )
                st.success("Opname toegevoegd aan de documentenlijst.")
            else:
                for document in st.session_state.document_list:
                    if document["name"] == name:
                        document["text"] = st.session_state.transcribed_text
                st.info("Bestaande opname bijgewerkt.")

            st.rerun()


def render_live_audit_log() -> None:
    if not st.session_state.live_audit_log:
        return

    st.divider()
    st.subheader("Sessielog")
    st.caption(f"{len(st.session_state.live_audit_log)} opname(s) in deze sessie")

    if st.button("Wis sessielog"):
        st.session_state.live_audit_log = []
        st.rerun()

    for index, entry in enumerate(reversed(st.session_state.live_audit_log), 1):
        log_number = len(st.session_state.live_audit_log) - index + 1
        with st.expander(f"Opname {log_number} — {entry['timestamp']}"):
            st.markdown(f"**Samenvatting:** {entry['summary']}")
            st.markdown(f"**Risicosignalen:** {len(entry['risk_results'])}")
            st.markdown(f"**ISO-bevindingen:** {len(entry['iso_data']['bevindingen'])}")

            report_data = genereer_live_audit_word_rapport(entry)
            st.download_button(
                "Download Word-rapport",
                data=report_data,
                file_name=f"live_audit_{entry['timestamp'].replace(':', '')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"dl_log_{index}",
            )


def genereer_live_audit_word_rapport(entry: dict) -> bytes:
    document = Document()
    document.add_heading(f"Live auditopname — {entry['timestamp']}", 0)

    document.add_heading("Transcriptie", level=1)
    document.add_paragraph(entry["transcriptie"])

    document.add_heading("Samenvatting", level=1)
    document.add_paragraph(entry["summary"])

    document.add_heading("Risicosignalen", level=1)
    for index, risk in enumerate(entry["risk_results"], 1):
        document.add_heading(
            f"{index}. {risk['issue_type'].upper()} "
            f"(score: {risk['score']:.2f})",
            level=2,
        )
        document.add_paragraph(risk["sentence"])

        if index <= len(entry["follow_up_questions"]):
            document.add_heading("Opvolgingsvragen", level=3)
            document.add_paragraph(entry["follow_up_questions"][index - 1])

    document.add_heading("ISO-bevindingen", level=1)
    for finding in entry["iso_data"]["bevindingen"]:
        document.add_heading(
            f"{finding['norm']} | {finding['clausule']}: {finding['titel']}",
            level=2,
        )
        document.add_paragraph(f"Ernst: {finding['ernst'].capitalize()}")
        document.add_paragraph(f"Probleem: {finding['beschrijving']}")
        document.add_paragraph(f"Aanbeveling: {finding['aanbeveling']}")

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def render_feedback_tab() -> None:
    st.subheader("Feedbackbeheer")
    st.info(
        "Feedback wordt opgeslagen in Supabase en wordt gebruikt als voorbeeld "
        "bij een volgende analyse."
    )

    feedback = st.session_state.feedback_store
    col1, col2 = st.columns(2)

    with col1:
        st.markdown(f"### Goede voorbeelden ({len(feedback['goede_vragen'])})")
        if not feedback["goede_vragen"]:
            st.info("Nog geen goede voorbeelden opgeslagen.")

        for index, entry in enumerate(reversed(feedback["goede_vragen"]), 1):
            with st.expander(f"{index}. {entry['issue_type'].upper()}"):
                st.markdown(f"**Zin:** {entry['zin']}")
                st.markdown(f"**Vragen:**\n{entry['vragen']}")
                if st.button("Verwijder", key=f"del_goed_{index}"):
                    verwijder_feedback_supabase(entry["zin"], entry["issue_type"])
                    st.session_state.feedback_store = laad_feedback_supabase()
                    st.rerun()

    with col2:
        st.markdown(f"### Minder passende voorbeelden ({len(feedback['slechte_vragen'])})")
        if not feedback["slechte_vragen"]:
            st.info("Nog geen minder passende voorbeelden opgeslagen.")

        for index, entry in enumerate(reversed(feedback["slechte_vragen"]), 1):
            with st.expander(f"{index}. {entry['issue_type'].upper()}"):
                st.markdown(f"**Zin:** {entry['zin']}")
                st.markdown(f"**Vragen:**\n{entry['vragen']}")
                if st.button("Verwijder", key=f"del_slecht_{index}"):
                    verwijder_feedback_supabase(entry["zin"], entry["issue_type"])
                    st.session_state.feedback_store = laad_feedback_supabase()
                    st.rerun()

    st.divider()
    if st.button("Wis alle feedback", type="secondary"):
        try:
            supabase_client.table("audit_feedback").delete().neq("id", "").execute()
            st.session_state.feedback_store = {"goede_vragen": [], "slechte_vragen": []}
            st.rerun()
        except Exception as exc:
            st.error(f"Fout: {exc}")

    st.divider()
    st.subheader("Werking van de feedbackloop")
    st.markdown(
        """
        1. Analyseer een document in de eerste tab.
        2. Beoordeel de gegenereerde vraagsets.
        3. Bij een volgende analyse worden de beoordelingen meegenomen als voorbeelden.
        4. De feedback blijft bewaard in Supabase.
        """
    )


def main() -> None:
    init_session_state()
    apply_css()
    render_header()

    (
        api_key,
        project_context,
        enable_9001,
        enable_14001,
        enable_45001,
        threshold,
        enable_audit_mode,
    ) = render_sidebar()

    tab1, tab2, tab3 = st.tabs(
        ["Document- en audioanalyse", "Live auditmodus", "Feedbackbeheer"]
    )

    with tab1:
        render_document_analysis_tab(
            api_key,
            project_context,
            enable_9001,
            enable_14001,
            enable_45001,
            threshold,
        )

    with tab2:
        render_live_audit_tab(
            api_key,
            enable_9001,
            enable_14001,
            enable_45001,
            enable_audit_mode,
        )

    with tab3:
        render_feedback_tab()


if __name__ == "__main__":
    main()
